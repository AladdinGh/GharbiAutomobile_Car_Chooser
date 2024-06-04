import logging
import pandas as pd
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx


# Import custom modules
from Helper_report_generation import preprocess_search_list , translate_preprocessed_search_result


# Setup logging configuration
logging.basicConfig(filename='processing.log', level=logging.INFO, 
                    format='%(asctime)s - %(levelname)s - %(message)s')

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
 
    #Add a hyperlink to a paragraph.
    
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add color if specified
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Add underline if specified
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Append the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

def add_features_table(doc, car, features):
  
    #Add a table with car features to the Word document.
    num_columns = 4  # Number of columns per row in the table
    table = doc.add_table(rows=1, cols=num_columns)
    table.style = 'Table Grid'

    row_cells = table.rows[0].cells
    col_idx = 0

    for feature in features:
        if feature not in ['Fahrzeughalter', 'Anzahl der Fahrzeughalter'] and feature in car.index and car[feature] == 1:
            row_cells[col_idx].text = feature
            col_idx += 1
            if col_idx == num_columns:
                row_cells = table.add_row().cells
                col_idx = 0

    # Add borders to the table cells
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

            tc_borders = OxmlElement('w:tcBorders')
            for key in ('left', 'right', 'top', 'bottom'):
                new_border = OxmlElement(f'w:{key}')
                new_border.set(qn('w:val'), 'single')
                tc_borders.append(new_border)
            cell._element.get_or_add_tcPr().append(tc_borders)

def print_report(file_path):
    try:
        # Read the data from the Excel file
        df = pd.read_excel(file_path)
        
        # Rank cars based on scores
        df_sorted = df.sort_values(by='Brutto Price', ascending=True)
        #best_fit_cars = df_sorted.head(10)
        best_fit_cars = df_sorted

        # Create a Word document
        doc = Document()

        # Generate the textual report
        doc.add_heading('Rapport Comparatif des voitures sur www.mobile.de', level=1)
        doc.add_paragraph(f"Nous avons trouvé {len(df)} voitures correspondant à vos critères.")
        doc.add_paragraph("\nLes 10 meilleures correspondances :")
        
        cpt = 0 
        for i, car in best_fit_cars.iterrows():
            cpt = cpt + 1 
            doc.add_paragraph(f"\n{car['Car Title']}", style='Title')
            doc.add_paragraph(f"\nNuméro de la voiture : {cpt}")
            #doc.add_paragraph(f"Score : {car['Score']:.2f}")
            doc.add_paragraph(f"Kilométrage : {car['Kilometerstand']} km")
            doc.add_paragraph(f"Date de mise en circulation : {car['Erstzulassung']}")
            doc.add_paragraph(f"Couleur : {car['Farbe'] if pd.notnull(car['Farbe']) else car['Farbe(constructeur)']}")
            doc.add_paragraph(f"Prix brut: {car['Brutto Price']} EUR")
            
            # Add clickable URL
            p = doc.add_paragraph("URL : ")
            add_hyperlink(p, car['Short_URL'], car['Short_URL'])

            doc.add_paragraph("Les options :")
            add_features_table(doc, car, df.columns)
            doc.add_paragraph("Description :")
            
            # Format description into paragraphs
            description = car['Description']
            paragraphs = re.split(r'[\n.]', description)
            for paragraph in paragraphs:
                doc.add_paragraph(f" {paragraph.strip()}")

        # Save the report to a Word document
        doc.save("Rapport comparatif.docx")
        logging.info("Report saved as 'Rapport comparatif.docx'")
    except Exception as e:
        logging.error(f"Error in print_report: {e}")

def prepare_and_print_report():
    try:
        
        # logging.info("preprocessing ....")
        # file_path = "search_list_car_features.xlsx" 
        # preprocess_search_list(file_path)
        
        
        
        # logging.info("Translating ....")
        # file_path = "preprocessed_search_result.xlsx"
        # translate_preprocessed_search_result(file_path)
        
        
        logging.info("Generating report ....")
        file_path = "translated_preprocessed_df.xlsx"
        print_report(file_path)
        
    except Exception as e:
        logging.error(f"Error in prepare_and_print_report: {e}")

prepare_and_print_report()