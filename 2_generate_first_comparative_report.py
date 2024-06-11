import logging
import pandas as pd
import re
from docx import Document

# Import custom modules
from Helper_generate_report import  add_hyperlink, add_features_table, assign_scores_report
from Helper_preprocess_translate import preprocess_search_list , translate_preprocessed_search_result

# Setup logging configuration
logging.basicConfig(filename='processing.log', level=logging.INFO, 
                    format='%(asctime)s - %(levelname)s - %(message)s')



def print_first_report(file_path):
    try:
        # Read the data from the Excel file
        df = pd.read_excel(file_path)
        
        # Rank cars based on scores
        df_sorted = df.sort_values(by='Score', ascending=False)
        best_fit_cars = df_sorted

        # Create a Word document
        doc = Document()

        # Generate the textual report
        doc.add_heading('Rapport Comparatif des voitures sur www.mobile.de', level=1)
        doc.add_paragraph(f"Nous avons trouvé {len(df)} voitures correspondant à vos critères.")
        
        for i, car in best_fit_cars.iterrows():
            doc.add_paragraph(f"\n{car['Car Title']}", style='Title')
            doc.add_paragraph(f"Numéro de la voiture : {car.name + 1}")
            doc.add_paragraph(f"Score : {car['Score']}")
            doc.add_paragraph(f"Kilométrage : {car['Kilometerstand']} km")
            doc.add_paragraph(f"Date de mise en circulation : {car['Erstzulassung']}")
            doc.add_paragraph(f"Couleur : {car['Farbe'] if pd.notnull(car['Farbe']) else car['Farbe(constructeur)']}")
            doc.add_paragraph(f"Prix brut: {car['Brutto Price']} EUR")
            
            # Add clickable URL
            p = doc.add_paragraph("URL : ")
            add_hyperlink(p, car['Short_URL'], car['Short_URL'])

            doc.add_paragraph("Les options :")
            add_features_table(doc, car, df.columns)
            doc.add_paragraph("Description :")
            
            # Format description into paragraphs
            description = str(car['Description']) if pd.notnull(car['Description']) else ""
            paragraphs = re.split(r'[\n.]', description)
            for paragraph in paragraphs:
                if paragraph.strip():  # Only add non-empty paragraphs
                    doc.add_paragraph(paragraph.strip())

        # Save the report to a Word document
        doc.save("output/Rapport_comparatif_initial.docx")
        logging.info("Report saved as 'Rapport_comparatif_initial.docx'")
    except Exception as e:
        logging.error(f"Error in print_first_report: {e}")

def prepare_and_print_first_report():
    try:
        
        flag_portal = "autoscout"
        logging.info("preprocessing ....")
        file_path = "autoscout_search_list_car_features.xlsx" 
        preprocess_search_list(file_path,flag_portal)
        
        
        logging.info("Translating ....")
        file_path = "output/1_preprocessed_search_result.xlsx"
        translate_preprocessed_search_result(file_path)
        
        flag_use_weights = False
        logging.info("Assigning score ....")
        file_path = "output/2_translated_preprocessed_df.xlsx"
        assign_scores_report(file_path,flag_use_weights)
        
        
        logging.info("Generating report ....")
        file_path = "output/3_translated_preprocessed_sorted_by_price_for_initial_report.xlsx"
        print_first_report(file_path)
        
    except Exception as e:
        logging.error(f"Error in prepare_and_print_report: {e}")

prepare_and_print_first_report()