import logging
import pandas as pd
from docx import Document
import re
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor



from Helper_generate_report import add_hyperlink, add_features_table, copy_rows_by_index, assign_scores_report, select_rows ,  add_features_table_light



def print_second_report(file_path):
    
    try:
        # Read the data from the Excel file
        df = pd.read_excel(file_path)
        
        # Rank cars based on scores
        df_sorted = df.sort_values(by='Score', ascending=False)
        #best_fit_cars = df_sorted.head(10)
        best_fit_cars = df_sorted

        # Create a Word document
        doc = Document()

        #####################################################################################################               
        # Generate the textual report
        doc.add_heading('Rapport Comparatif des voitures sélectionnées', level=1)
        doc.add_paragraph(f"Les {len(df)} voitures correspondant à vos critères sont triées par score")
        
        
         # Add the table with the specified columns
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        # Define the headers
        headers = ["Score", "Titre", "Prix brut", "Kilométrage", "Date Circulation", "Nombre d'options", "Lien"]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black color
            hdr_cells[i].paragraphs[0].runs[0].font.name = 'Calibri'  # Set font name
        
        # Add data to the table
        for idx, car in best_fit_cars.iterrows():
            row_cells = table.add_row().cells
            #row_cells[0].text = str(car['index']+1)  # Adjusting index to match your report
            row_cells[0].text = str(round(car['Score'], 2))  # Round score to 2 decimal places
            row_cells[1].text = car['Car Title']
            row_cells[2].text = f"{car['Brutto Price']}"
            row_cells[3].text = f"{car['Kilometerstand']}"
            row_cells[4].text = str(car['Erstzulassung'])  # Convert to string to avoid TypeError
            row_cells[5].text = str(add_features_table_light(car, df.columns))   # Count number of options
            
            # Add hyperlink to cell if URL is not empty
            if car['Short_URL']:
                add_hyperlink(row_cells[6].paragraphs[0], car['Short_URL'], car['Short_URL'])
            
            # Apply font and style to data cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.name = 'Calibri'  # Set font name
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set cell widths
        widths = [0.5, 1, 2, 1.5, 1.5, 1.5, 1]
        for i, width in enumerate(widths):
            table.columns[i].width = Pt(width * 72)  # Convert inches to points (1 inch = 72 points)
        
        # Apply borders to the table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                        
###############################################################################################################
        doc.add_heading("Description détaillée de votre sélection", level=2)
        for i, car in best_fit_cars.iterrows():
            #cpt = cpt + 1 
            doc.add_paragraph(f"\n{car['Car Title']}", style='Title')
            doc.add_paragraph(f"Numéro de la voiture : {car['index']+1}")
            doc.add_paragraph(f"Puissance : {car['KW']} KW")
            doc.add_paragraph(f"Consommation : {car['Kraftstoffverbrauch']} L/100km")
            doc.add_paragraph(f"Boite de vitesse : {car['Getriebe']}")
            #doc.add_paragraph(f"carburant: {car['Kraftstoffart']}")
            doc.add_paragraph(f"Couleur : {car['Farbe'] if pd.notnull(car['Farbe']) else car['Farbe(constructeur)']}")
            doc.add_paragraph(f"Prix brut: {car['Brutto Price']} EUR")
            
            
            # Add clickable URL
            # p = doc.add_paragraph("URL : ")
            # add_hyperlink(p, car['Short_URL'], car['Short_URL'])

            doc.add_paragraph("Les options :")
            add_features_table(doc, car, df.columns)
            doc.add_paragraph("Description :")
            
            # Format description into paragraph
            description = str(car['Description'])
            paragraphs = re.split(r'[\n.]', description)
            for paragraph in paragraphs:
                doc.add_paragraph(f" {paragraph.strip()}")

        # Save the report to a Word document
        doc.save("output/Rapport_comparatif_final.docx")
        logging.info("Report saved as 'Rapport_comparatif_final.docx'")
    except Exception as e:
        logging.error(f"Error in print_report: {e}")




def prepare_and_print_second_report():
    try:

        # get only the columns from the client
        select_rows()
        
        file_path = "output/selected_rows_df.xlsx"
        logging.info("Assigning score ....")
        assign_scores_report(file_path,True)
        
        
        logging.info("Generating report ....")
        file_path = "output/3_translated_preprocessed_sorted_by_score_for_second_report.xlsx"
        print_second_report(file_path)
        
    except Exception as e:
        logging.error(f"Error in prepare_and_print_report: {e}")

prepare_and_print_second_report()